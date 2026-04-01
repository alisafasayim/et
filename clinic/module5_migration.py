"""
Modül 5: Samsung Notes → Notion Toplu Migrasyon

Kurulum:
    pip install python-docx markdown-it-py requests

Çevre değişkenleri:
    NOTION_TOKEN         - Notion Integration secret (secret_xxx)
    NOTION_DATABASE_ID   - Hedef veritabanı ID

Kullanım:
    python module5_migration.py --dir ./samsung_notes
    python module5_migration.py --dir ./samsung_notes --ext md
    python module5_migration.py --dir ./samsung_notes --dry-run
"""

import argparse
import json
import logging
import os
import re
import time
from pathlib import Path

import requests
from docx import Document
from markdown_it import MarkdownIt

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("migration")

# ---------------------------------------------------------------------------
# Sabitler
# ---------------------------------------------------------------------------

NOTION_TOKEN = os.getenv("NOTION_TOKEN", "")
NOTION_DATABASE_ID = os.getenv("NOTION_DATABASE_ID", "")
NOTION_API_VERSION = "2022-06-28"
NOTION_BASE_URL = "https://api.notion.com/v1"

# Notion tek istekte max 100 block kabul eder
NOTION_BLOCK_CHUNK_SIZE = 100

# Hız sınırı: Notion API saniyede ~3 istek (ücretsiz workspace)
REQUEST_DELAY_SEC = 0.4

# ---------------------------------------------------------------------------
# 1. Notion HTTP İstemcisi
# ---------------------------------------------------------------------------

def _headers() -> dict:
    if not NOTION_TOKEN:
        raise EnvironmentError("NOTION_TOKEN çevre değişkeni ayarlanmamış.")
    return {
        "Authorization": f"Bearer {NOTION_TOKEN}",
        "Notion-Version": NOTION_API_VERSION,
        "Content-Type": "application/json",
    }


def _post(endpoint: str, payload: dict) -> dict:
    url = f"{NOTION_BASE_URL}{endpoint}"
    resp = requests.post(url, headers=_headers(), json=payload, timeout=30)
    resp.raise_for_status()
    time.sleep(REQUEST_DELAY_SEC)
    return resp.json()


def _patch(endpoint: str, payload: dict) -> dict:
    url = f"{NOTION_BASE_URL}{endpoint}"
    resp = requests.patch(url, headers=_headers(), json=payload, timeout=30)
    resp.raise_for_status()
    time.sleep(REQUEST_DELAY_SEC)
    return resp.json()


# ---------------------------------------------------------------------------
# 2. Notion Block Yardımcıları
# ---------------------------------------------------------------------------

def _rich_text(text: str) -> list:
    """2000 karakter sınırını aşan metni bölerek rich_text listesi döner."""
    chunks = [text[i:i+2000] for i in range(0, len(text), 2000)] if text else [""]
    return [{"type": "text", "text": {"content": chunk}} for chunk in chunks]


def _paragraph(text: str) -> dict:
    return {
        "object": "block", "type": "paragraph",
        "paragraph": {"rich_text": _rich_text(text)},
    }


def _heading1(text: str) -> dict:
    return {
        "object": "block", "type": "heading_1",
        "heading_1": {"rich_text": _rich_text(text[:2000])},
    }


def _heading2(text: str) -> dict:
    return {
        "object": "block", "type": "heading_2",
        "heading_2": {"rich_text": _rich_text(text[:2000])},
    }


def _heading3(text: str) -> dict:
    return {
        "object": "block", "type": "heading_3",
        "heading_3": {"rich_text": _rich_text(text[:2000])},
    }


def _bulleted_item(text: str) -> dict:
    return {
        "object": "block", "type": "bulleted_list_item",
        "bulleted_list_item": {"rich_text": _rich_text(text[:2000])},
    }


def _numbered_item(text: str) -> dict:
    return {
        "object": "block", "type": "numbered_list_item",
        "numbered_list_item": {"rich_text": _rich_text(text[:2000])},
    }


def _divider() -> dict:
    return {"object": "block", "type": "divider", "divider": {}}


def _code_block(text: str, language: str = "plain text") -> dict:
    return {
        "object": "block", "type": "code",
        "code": {
            "rich_text": _rich_text(text[:2000]),
            "language": language,
        },
    }


def _quote(text: str) -> dict:
    return {
        "object": "block", "type": "quote",
        "quote": {"rich_text": _rich_text(text[:2000])},
    }


def _append_blocks(page_id: str, blocks: list[dict], dry_run: bool = False) -> None:
    """100'lük chunk'larla Notion'a block ekler."""
    for i in range(0, len(blocks), NOTION_BLOCK_CHUNK_SIZE):
        chunk = blocks[i:i + NOTION_BLOCK_CHUNK_SIZE]
        if dry_run:
            logger.debug("  [dry-run] %d block append edilecek", len(chunk))
            continue
        _patch(f"/blocks/{page_id}/children", {"children": chunk})


# ---------------------------------------------------------------------------
# 3. .docx Parser
# ---------------------------------------------------------------------------

def parse_docx(file_path: Path) -> list[dict]:
    """
    python-docx ile .docx dosyasını okur.
    Başlık stillerini (Heading 1/2/3) ve paragrafları Notion bloklarına dönüştürür.
    """
    doc = Document(str(file_path))
    blocks: list[dict] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name.lower()

        if "heading 1" in style:
            blocks.append(_heading1(text))
        elif "heading 2" in style:
            blocks.append(_heading2(text))
        elif "heading 3" in style:
            blocks.append(_heading3(text))
        elif "list bullet" in style:
            blocks.append(_bulleted_item(text))
        elif "list number" in style:
            blocks.append(_numbered_item(text))
        else:
            blocks.append(_paragraph(text))

    # Tablolar
    for table in doc.tables:
        blocks.append(_divider())
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                blocks.append(_paragraph(row_text))
        blocks.append(_divider())

    return blocks


# ---------------------------------------------------------------------------
# 4. Markdown Parser
# ---------------------------------------------------------------------------

def parse_markdown(file_path: Path) -> list[dict]:
    """
    Samsung Notes'tan dışa aktarılan Markdown dosyasını satır satır okur.
    CommonMark uyumlu yapıları Notion bloklarına dönüştürür.
    """
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    lines = raw.splitlines()
    blocks: list[dict] = []

    i = 0
    in_code_block = False
    code_lines: list[str] = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # --- Kod bloğu ---
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip() or "plain text"
                code_lines = []
            else:
                in_code_block = False
                blocks.append(_code_block("\n".join(code_lines), code_lang))
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Başlıklar ---
        if line.startswith("### "):
            blocks.append(_heading3(line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(_heading2(line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(_heading1(line[2:].strip()))

        # --- Yatay çizgi ---
        elif re.match(r"^[-*_]{3,}$", line.strip()):
            blocks.append(_divider())

        # --- Alıntı ---
        elif line.startswith("> "):
            blocks.append(_quote(line[2:].strip()))

        # --- Madde işaretli liste ---
        elif re.match(r"^[-*+] ", line):
            blocks.append(_bulleted_item(line[2:].strip()))

        # --- Numaralı liste ---
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line).strip()
            blocks.append(_numbered_item(text))

        # --- Boş satır → atla ---
        elif line.strip() == "":
            pass

        # --- Düz metin ---
        else:
            # Markdown inline formatlamayı sade metne dönüştür
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)   # bold
            text = re.sub(r"\*(.+?)\*", r"\1", text)        # italic
            text = re.sub(r"`(.+?)`", r"\1", text)          # inline code
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text) # link → sadece metin
            blocks.append(_paragraph(text.strip()))

        i += 1

    return blocks


# ---------------------------------------------------------------------------
# 5. Notion Sayfası Oluşturma
# ---------------------------------------------------------------------------

def create_notion_page(title: str, dry_run: bool = False) -> str:
    """
    Notion veritabanında başlıkla yeni bir sayfa oluşturur.
    Döner: page_id
    """
    if not NOTION_DATABASE_ID:
        raise EnvironmentError("NOTION_DATABASE_ID çevre değişkeni ayarlanmamış.")

    payload = {
        "parent": {"database_id": NOTION_DATABASE_ID},
        "properties": {
            "Hasta Adı": {
                "title": [{"text": {"content": title[:2000]}}]
            }
        },
    }

    if dry_run:
        logger.info("  [dry-run] Sayfa oluşturulacak: '%s'", title)
        return "dry-run-page-id"

    result = _post("/pages", payload)
    page_id = result["id"]
    logger.info("  Sayfa oluşturuldu: '%s' → %s", title, page_id)
    return page_id


# ---------------------------------------------------------------------------
# 6. Tek Dosya Migrasyonu
# ---------------------------------------------------------------------------

def migrate_file(file_path: Path, dry_run: bool = False) -> dict:
    """
    Tek bir .docx veya .md dosyasını Notion'a aktarır.
    Döner: {'file', 'title', 'page_id', 'block_count', 'status'}
    """
    title = file_path.stem  # Dosya adı (uzantısız) → sayfa başlığı
    ext = file_path.suffix.lower()

    logger.info("Aktarılıyor: %s", file_path.name)

    try:
        if ext == ".docx":
            blocks = parse_docx(file_path)
        elif ext in (".md", ".markdown"):
            blocks = parse_markdown(file_path)
        else:
            return {"file": str(file_path), "status": "skipped", "reason": "desteklenmeyen format"}

        if not blocks:
            logger.warning("  İçerik bulunamadı: %s", file_path.name)
            return {"file": str(file_path), "title": title, "status": "empty"}

        # Sayfa oluştur
        page_id = create_notion_page(title, dry_run=dry_run)

        # Blokları ekle
        _append_blocks(page_id, blocks, dry_run=dry_run)

        logger.info("  Tamamlandı: %d block | page_id: %s", len(blocks), page_id)
        return {
            "file": file_path.name,
            "title": title,
            "page_id": page_id,
            "block_count": len(blocks),
            "status": "ok",
        }

    except Exception as exc:
        logger.error("  HATA [%s]: %s", file_path.name, exc)
        return {"file": file_path.name, "title": title, "status": "failed", "error": str(exc)}


# ---------------------------------------------------------------------------
# 7. Toplu Migrasyon Orkestratörü
# ---------------------------------------------------------------------------

def migrate_directory(
    source_dir: Path,
    extensions: list[str] | None = None,
    dry_run: bool = False,
) -> list[dict]:
    """
    Belirli bir klasördeki tüm .docx/.md dosyalarını sırayla Notion'a aktarır.

    source_dir  : Samsung Notes dışa aktarım klasörü
    extensions  : ['docx', 'md'] gibi filtre; None ise her ikisi de işlenir
    dry_run     : True ise Notion'a hiçbir şey yazılmaz, yalnızca loglanır
    """
    if not source_dir.exists():
        raise FileNotFoundError(f"Kaynak klasör bulunamadı: {source_dir}")

    allowed_exts = {f".{e.lstrip('.')}" for e in (extensions or ["docx", "md", "markdown"])}
    files = sorted(
        f for f in source_dir.rglob("*")
        if f.is_file() and f.suffix.lower() in allowed_exts
    )

    if not files:
        logger.warning("İşlenecek dosya bulunamadı: %s (uzantılar: %s)", source_dir, allowed_exts)
        return []

    logger.info(
        "Migrasyon başlıyor | %d dosya | dry_run=%s | hedef DB: %s",
        len(files), dry_run, NOTION_DATABASE_ID[:8] + "..." if NOTION_DATABASE_ID else "—",
    )

    results = []
    for idx, file_path in enumerate(files, start=1):
        logger.info("[%d/%d] %s", idx, len(files), file_path.name)
        result = migrate_file(file_path, dry_run=dry_run)
        results.append(result)

    # Özet
    ok = sum(1 for r in results if r["status"] == "ok")
    failed = sum(1 for r in results if r["status"] == "failed")
    skipped = sum(1 for r in results if r["status"] in ("skipped", "empty"))

    logger.info(
        "Migrasyon tamamlandı | Başarılı: %d | Başarısız: %d | Atlandı: %d",
        ok, failed, skipped,
    )

    # Başarısız dosyaları raporla
    report_path = source_dir / "migration_report.json"
    if not dry_run:
        report_path.write_text(
            json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        logger.info("Rapor kaydedildi: %s", report_path)

    return results


# ---------------------------------------------------------------------------
# 8. CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Samsung Notes → Notion Toplu Migrasyon",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Örnekler:
  python module5_migration.py --dir ./samsung_notes
  python module5_migration.py --dir ./samsung_notes --ext md
  python module5_migration.py --dir ./samsung_notes --ext docx --dry-run
        """,
    )
    parser.add_argument(
        "--dir", required=True, type=Path,
        help="Samsung Notes dışa aktarım klasörü",
    )
    parser.add_argument(
        "--ext", choices=["docx", "md", "both"], default="both",
        help="İşlenecek dosya uzantısı (varsayılan: both)",
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Notion'a yazmadan simülasyon yap",
    )
    args = parser.parse_args()

    extensions = (
        ["docx", "md"] if args.ext == "both" else [args.ext]
    )

    migrate_directory(
        source_dir=args.dir,
        extensions=extensions,
        dry_run=args.dry_run,
    )
